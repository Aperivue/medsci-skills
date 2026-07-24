#!/usr/bin/env python3
"""Response-letter claim <-> revised-manuscript verification (rule-backed gate).

A response-to-reviewers letter asserts concrete edits: "we added the sentence
'...'", "we now cite Tariq et al. [15]". The single-source-of-truth is the
*revised manuscript*, not the response prose — yet a claimed edit can be absent
from the body (a real incident: an added Discussion citation was described in the
response but never inserted, and both a reviewer round and the authors missed it
until a body grep). This gate makes that class deterministic for both sides:
`/revise` (author, before sending) and `/peer-review` (reviewer, verifying the
author's claims against the revised manuscript).

It is deliberately conservative — it verifies only claims carrying a strong,
checkable anchor, so paraphrase and honest rewording do not false-positive:

  * RESPONSE_QUOTE_UNVERIFIED (major) — the letter says specific text was
    added / inserted / "now reads" and quotes it verbatim, but that quoted text
    is absent from the revised manuscript body.
  * RESPONSE_QUOTE_UNRESOLVED (minor) — the quoted text IS there in order, but
    only once foreign tokens are allowed between its words, or a word or two is
    missing. That is the signature of a dirty extraction (a bled reference
    column, PDF line numbers, a footnote marker, a hyphen split across a line),
    not of an edit that was never made. Reported so a human looks; never counted
    as drift. A contiguous substring test cannot tell these apart and calls a
    correct quote absent — the failure that once came one step from having two
    accurate verbatim quotes deleted. Matching lives in _quote_match.py.
  * RESPONSE_CITATION_UNVERIFIED (major) — the letter says a citation was added
    / "now cite(d)", but none of the cited tokens ([N] / [@key] / Author et al.)
    appear in the revised manuscript body.

Vague claims with no quote and no citation ("we clarified the Methods") are not
verifiable and are intentionally NOT flagged. Reviewer-comment blockquotes
(lines beginning with '>') are excluded so the reviewer's own quoted text is
never mistaken for an author addition.

Usage:
    check_response_claims.py --response response.md --manuscript revised.md [--strict]
    check_response_claims.py --response r.md --manuscript revised.docx --out qc/response_claims.json

Exit 0 when every anchored claim is verified (or none exist). With --strict,
exit 1 if any major verdict fires. Stdlib only; .docx read via python-docx when
available.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _quote_match import match_quality  # noqa: E402  (same-dir helper)

# A claim that asserts an addition/edit to the manuscript.
CLAIM_VERB = re.compile(
    r"\b("
    r"added the (?:sentence|statement|clause|text|following|phrase)|"
    r"added a (?:sentence|statement|clause|citation|reference|paragraph)|"
    r"we (?:have )?added|have added|now added|"
    r"inserted|included the (?:sentence|statement|text|citation|reference)|"
    r"now (?:reads|read|states|state)|"
    r"(?:revised|changed|reworded|rephrased|amended) [^.\n]{0,60}? to (?:read|state)|"
    r"now cites?|now cited|we (?:now )?cite|added (?:the )?(?:citation|reference)s?"
    r")\b",
    re.IGNORECASE,
)

# Quoted string: straight or curly, >= 12 chars (a sentence-like assertion).
QUOTE = re.compile(r"[\"“‘']([^\"“”‘’']{12,600})[\"”’']")

# Citation tokens claimed as added.
CIT_NUMERIC = re.compile(r"\[(\d{1,3}(?:\s*[,–-]\s*\d{1,3})*)\]")
CIT_BIBKEY = re.compile(r"\[@([A-Za-z0-9_:.\-]+)\]")
CIT_AUTHOR = re.compile(r"\b([A-Z][A-Za-zÀ-ſ'-]{2,})\s+et\s+al\.?")

WINDOW = 320  # chars after a claim verb to look for its object


def read_text(path: Path) -> str:
    """Return plain text from .md/.txt or .docx (recursive paragraphs + tables)."""
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document  # type: ignore
            from docx.document import Document as _Doc  # noqa: F401
        except Exception as exc:  # pragma: no cover
            raise SystemExit(f"python-docx required to read {path}: {exc}")
        doc = Document(str(path))
        parts: list[str] = []

        def walk_table(tbl):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
                    for t in cell.tables:
                        walk_table(t)

        for p in doc.paragraphs:
            parts.append(p.text)
        for t in doc.tables:
            walk_table(t)
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def normalize(s: str) -> str:
    """Casefold + collapse whitespace + strip markdown emphasis for substring match."""
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("’", "'").replace("‘", "'")
    s = s.replace("“", '"').replace("”", '"')
    s = re.sub(r"[*_`]", "", s)  # markdown emphasis / code ticks
    s = re.sub(r"\s+", " ", s)
    return s.casefold().strip()


def strip_response_blockquotes(text: str) -> str:
    """Drop reviewer-comment blockquote lines (>) so their quotes aren't scanned."""
    keep = [ln for ln in text.splitlines() if not ln.lstrip().startswith(">")]
    return "\n".join(keep)


def extract_claims(response: str):
    """Yield (kind, anchor, context) for anchored addition claims in Response prose."""
    prose = strip_response_blockquotes(response)
    claims = []
    for m in CLAIM_VERB.finditer(prose):
        start = m.start()
        window = prose[start : start + WINDOW]
        ctx = re.sub(r"\s+", " ", prose[max(0, start - 20) : start + 120]).strip()
        # quoted additions
        for q in QUOTE.finditer(window):
            text = q.group(1).strip()
            if len(text.split()) >= 4:
                claims.append(("quote", text, ctx))
        # citation additions
        cits = []
        for cm in CIT_NUMERIC.finditer(window):
            for n in re.split(r"[,–-]", cm.group(1)):
                if n.strip():
                    cits.append(("num", n.strip()))
        for cm in CIT_BIBKEY.finditer(window):
            cits.append(("key", cm.group(1)))
        for cm in CIT_AUTHOR.finditer(window):
            cits.append(("author", cm.group(1)))
        # only treat a verb as a citation claim if the verb itself is citation-ish
        if cits and re.search(r"cit|reference", m.group(0), re.IGNORECASE):
            claims.append(("citation", cits, ctx))
    return claims


def grade_quote(body: str, quote: str) -> dict:
    """Grade the quote's presence in the body via the extraction-tolerant matcher.

    A contiguous search alone is not safe here: the manuscript may arrive as a .docx whose
    extraction wedges a footnote marker, a line number, or a bled column of reference text
    into the middle of the very sentence being checked. Those quotes are present and correct,
    and a substring test calls them absent — the failure that once nearly had two accurate
    verbatim quotes deleted. See _quote_match.py."""
    return match_quality(quote, body)


def body_has_citation(body: str, norm_body: str, cits) -> bool:
    """True if ANY cited token appears in the body (conservative: any-match passes)."""
    for kind, tok in cits:
        if kind == "num" and re.search(r"\[\s*\d*[,\s–-]*" + re.escape(tok) + r"\b", body):
            return True
        if kind == "num" and ("[" + tok + "]") in body:
            return True
        if kind == "key" and ("@" + tok) in body:
            return True
        if kind == "author" and normalize(tok) in norm_body:
            return True
    return False


def build_report(response_path: Path, manuscript_path: Path) -> dict:
    response = read_text(response_path)
    body = read_text(manuscript_path)
    norm_body = normalize(body)
    findings = []
    for kind, anchor, ctx in extract_claims(response):
        if kind == "quote":
            g = grade_quote(body, anchor)
            if g["grade"] == "INTERLEAVED":
                findings.append(
                    {
                        "verdict": "RESPONSE_QUOTE_UNRESOLVED",
                        "severity": "minor",
                        "claimed_text": anchor,
                        "context": ctx,
                        "match": g,
                        "message": (
                            f"Every word of the quoted text appears in order, but with {g['inserted']} "
                            "foreign token(s) wedged in — consistent with a dirty extraction (a bled "
                            "reference column, line numbers, a footnote marker), not a missing edit. "
                            "Confirm by eye; do not treat as absent."
                        ),
                    }
                )
            elif g["grade"] == "PARTIAL":
                findings.append(
                    {
                        "verdict": "RESPONSE_QUOTE_UNRESOLVED",
                        "severity": "minor",
                        "claimed_text": anchor,
                        "context": ctx,
                        "match": g,
                        "message": (
                            f"{g['matched']} of {g['total']} words of the quoted text appear in order "
                            "— enough to be the same sentence damaged in extraction (a hyphen split "
                            "across a line, a dropped glyph) rather than an edit that was never made. "
                            "Confirm by eye."
                        ),
                    }
                )
            elif g["grade"] == "ABSENT":
                findings.append(
                    {
                        "verdict": "RESPONSE_QUOTE_UNVERIFIED",
                        "severity": "major",
                        "claimed_text": anchor,
                        "context": ctx,
                        "match": g,
                        "message": "Response quotes added text that is absent from the revised manuscript body.",
                    }
                )
        elif kind == "citation":
            if not body_has_citation(body, norm_body, anchor):
                findings.append(
                    {
                        "verdict": "RESPONSE_CITATION_UNVERIFIED",
                        "severity": "major",
                        "claimed_citation": [t for _, t in anchor],
                        "context": ctx,
                        "message": "Response claims a citation was added but none of the cited tokens appear in the revised manuscript body.",
                    }
                )
    n_major = sum(1 for f in findings if f["severity"] == "major")
    return {
        "response": str(response_path),
        "manuscript": str(manuscript_path),
        "findings": findings,
        "summary": {"major": n_major, "unresolved": len(findings) - n_major},
        # An UNRESOLVED quote is a "look at this", not a defect: the words are demonstrably
        # there and only the extraction is suspect. Safety therefore turns on MAJOR findings,
        # which is what --strict has always documented.
        "submission_safe": n_major == 0,
    }


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--response", required=True, type=Path, help="response-to-reviewers letter (.md/.txt/.docx)")
    ap.add_argument("--manuscript", required=True, type=Path, help="revised manuscript (.md/.txt/.docx)")
    ap.add_argument("--out", type=Path, help="write JSON report here")
    ap.add_argument("--strict", action="store_true", help="exit 1 if any major verdict fires")
    ap.add_argument("--quiet", action="store_true")
    args = ap.parse_args(argv)

    for p in (args.response, args.manuscript):
        if not p.is_file():
            print(f"error: file not found: {p}", file=sys.stderr)
            return 2

    report = build_report(args.response, args.manuscript)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(json.dumps({"detector": "check_response_claims", **report}, indent=2, ensure_ascii=False), encoding="utf-8")

    if not args.quiet:
        s = report["summary"]
        if not report["findings"]:
            print("OK: every anchored response claim is verified against the revised manuscript.")
        else:
            print(f"RESPONSE_CLAIM findings — {s['major']} major, {s['unresolved']} unresolved:")
            for f in report["findings"]:
                anchor = f.get("claimed_text") or ", ".join(f.get("claimed_citation", []))
                print(f"  [{f['verdict']}] ({f['severity']}) {anchor!r}")
                print(f"      near: {f['context']}")
                if f["severity"] == "minor":
                    print(f"      {f['message']}")
            if s["major"] == 0:
                print("\nNo major drift: the unresolved item(s) are extraction-quality doubts, "
                      "not claims of an edit that was never made.")

    if args.strict and not report["submission_safe"]:
        print("\nRESPONSE_CLAIM_UNVERIFIED: a response-letter claim is not reflected in the revised manuscript.", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
