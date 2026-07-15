#!/usr/bin/env python3
""""Your paper is too dense" is the one comment you cannot address by adding text.

A real revision, measured straight out of the .docx files of a DTA meta-analysis:

    v18  submitted       7,172 words  ->  four reviewers: "too dense / shorten / move to supplement"
    v20  THE revision    7,785 words  ->  +613.  Every named term went UP (kappa 5->6, Deeks 10->11)
         answering those                    because each comment was answered point-by-point, and
         comments                           point-by-point response REWARDS adding text.
    v21  proof revision  6,439 words  ->  -1,346 (733 below the original).  <- the accepted version.

Answering "your text is too dense" comment-by-comment made it denser. It took three rounds to do
what four reviewers asked in round one, because point-by-point culture rewards showing you addressed
each comment — and length is the one comment adding text cannot address.

This gate is pure arithmetic. If the reviewer comments contain a density/length complaint AND the
revised manuscript body did not get SHORTER than the previous version, the complaint was not
addressed — it was made worse. It fires on v20 immediately, and stays silent on v21.

It reads:
  --comments   the reviewer decision letter (where the complaint lives)
  --previous   the manuscript as the reviewers saw it   (the word count they complained about)
  --revised    the manuscript you are about to send back (must be shorter, if a complaint was made)

Body word count is measured the way a reviewer counts it: Introduction through Discussion, excluding
the title page, abstract, references, tables, figure legends, and — so a citation-heavy paragraph is
not scored as verbose — bracketed/parenthetical citation markers. It is a ratio check, so the exact
rule matters less than applying it identically to both versions, which it does.

Usage:
    check_density_complaint.py --comments letter.md --previous v_prev.docx --revised v_new.docx \
        [--out qc/density.json] [--strict]

Stdlib only (.docx via python-docx when present).
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

# The vocabulary a reviewer reaches for when a manuscript is too long or too dense. Each is a
# complaint that adding text cannot answer.
COMPLAINT = re.compile(
    r"\b(too (long|dense|detailed|verbose|wordy)|overly (long|detailed|dense|complex)|"
    r"shorten(ed|ing)?|condens(e|ed|ing)|trim(med|ming)?|cut down|tighten(ed|ing)?|"
    r"reduce (the )?(length|word count|detail)|"
    r"(mov\w+|belongs?).{0,30}?(to (the )?supplement|to supplementary|in (the )?supplement)|"
    r"excessive(ly)? (detail|length|long)|"
    r"(somewhat |very |quite |rather |overly )?(dense|verbose|wordy)\b|"
    r"(is|are|reads?|too) (very )?(repetitive|redundant)|"
    r"difficult to follow|hard to follow|dilut(e|es|ing) the (main )?message|"
    r"streamlin(e|ed|ing)|language editing|more concise|less dense|for (legibility|readability)|"
    r"repeat(ed|s|ing)? (multiple times|throughout|itself))\b",
    re.IGNORECASE,
)

# Section headings, in reviewer-count order. Body = first Introduction..Introduction-like heading
# through the end of Discussion / before References.
BODY_START = re.compile(r"^#{0,4}\s*\**\s*(introduction|background)\b", re.IGNORECASE | re.MULTILINE)
BODY_END = re.compile(r"^#{0,4}\s*\**\s*(references|acknowledg|funding|conflict|"
                      r"data availability|supplementary|supporting information|figure legends?)\b",
                      re.IGNORECASE | re.MULTILINE)
ABSTRACT = re.compile(r"^#{0,4}\s*\**\s*abstract\b", re.IGNORECASE | re.MULTILINE)

CITATION = re.compile(r"\[[\d,\s–\-]+\]|\((?:[A-Z][A-Za-z'`-]+(?: et al\.?)?,?\s*\d{4}[a-z]?;?\s*)+\)")


def read_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise SystemExit(f"python-docx required to read {path}: {exc}")
        doc = Document(str(path))
        parts: list[str] = []

        def walk(tbl):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
                    for t in cell.tables:
                        walk(t)

        for p in doc.paragraphs:
            parts.append(p.text)
        for t in doc.tables:
            walk(t)
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def body_word_count(text: str) -> int:
    """Introduction..before-References, abstract excluded, citation markers removed."""
    text = unicodedata.normalize("NFKC", text)
    m0 = BODY_START.search(text)
    body = text[m0.start():] if m0 else text
    m1 = BODY_END.search(body)
    if m1:
        body = body[: m1.start()]
    # if an Abstract heading is inside our slice (no Introduction found), drop it
    if not m0:
        a = ABSTRACT.search(body)
        if a:
            nxt = BODY_START.search(body, a.end())
            if nxt:
                body = body[nxt.start():]
    body = CITATION.sub(" ", body)
    body = re.sub(r"^#{1,6}\s.*$", " ", body, flags=re.MULTILINE)  # drop heading lines themselves
    body = re.sub(r"[*_`>|#]", " ", body)                          # markdown punctuation
    words = re.findall(r"[A-Za-z0-9][\w'–-]*", body)
    return len(words)


def complaints(text: str) -> list[str]:
    """The density/length complaint sentences in the decision letter."""
    out: list[str] = []
    for sent in re.split(r"(?<=[.!?])\s+|\n", text):
        if COMPLAINT.search(sent):
            out.append(re.sub(r"\s+", " ", sent).strip()[:160])
    # dedupe preserving order
    seen, uniq = set(), []
    for s in out:
        if s.lower() not in seen:
            seen.add(s.lower()); uniq.append(s)
    return uniq


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[0])
    ap.add_argument("--comments", required=True, type=Path, help="reviewer decision letter (.md/.txt/.docx)")
    ap.add_argument("--previous", required=True, type=Path, help="manuscript the reviewers saw (.md/.txt/.docx)")
    ap.add_argument("--revised", required=True, type=Path, help="manuscript you are about to send back")
    ap.add_argument("--out", type=Path, help="write JSON report here")
    ap.add_argument("--strict", action="store_true", help="exit 1 if the complaint was not addressed")
    ap.add_argument("--quiet", action="store_true")
    a = ap.parse_args(argv)

    for p in (a.comments, a.previous, a.revised):
        if not p.is_file():
            print(f"error: file not found: {p}", file=sys.stderr)
            return 2

    comp = complaints(read_text(a.comments))
    prev_wc = body_word_count(read_text(a.previous))
    new_wc = body_word_count(read_text(a.revised))
    delta = new_wc - prev_wc

    fired = bool(comp) and delta >= 0
    report = {
        "detector": "check_density_complaint",
        "density_complaints": comp,
        "previous_body_words": prev_wc,
        "revised_body_words": new_wc,
        "delta_words": delta,
        "verdict": "DENSITY_COMPLAINT_UNADDRESSED" if fired else "OK",
    }
    if a.out:
        a.out.parent.mkdir(parents=True, exist_ok=True)
        a.out.write_text(json.dumps(report, indent=2), encoding="utf-8")

    if not a.quiet:
        if not comp:
            print(f"OK: no density/length complaint in the decision letter "
                  f"(body {prev_wc} -> {new_wc} words, {delta:+d}).")
        elif not fired:
            print(f"OK: a density complaint was raised and the body got shorter "
                  f"({prev_wc} -> {new_wc} words, {delta:+d}).")
        else:
            print(f"DENSITY_COMPLAINT_UNADDRESSED: reviewers said the manuscript was too dense/long, "
                  f"and the body did not get shorter ({prev_wc} -> {new_wc} words, {delta:+d}).\n")
            for c in comp[:6]:
                print(f"  reviewer: {c}")
            print(
                "\n'Too dense' is the one comment you cannot address by adding text, and point-by-point\n"
                "response rewards adding it. Answering each density comment individually made the last\n"
                "manuscript that hit this LONGER by 613 words; the accepted version was 733 words below\n"
                "where it started. Cut, or move detail to the supplement — do not defend length by\n"
                "adding a paragraph that explains it."
            )
    return 1 if (fired and a.strict) else 0


if __name__ == "__main__":
    sys.exit(main())
